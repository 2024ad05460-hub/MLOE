from __future__ import annotations
import csv
import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
FIG = REPORTS / "figures"

m1 = json.loads((ROOT / "training/models/m1_metrics.json").read_text())
m2 = json.loads((ROOT / "training/models/m2_metrics.json").read_text())
m3 = json.loads((ROOT / "training/models/m3_metrics.json").read_text())
bench = json.loads((ROOT / "optimisation/results/benchmark_results.json").read_text())["results"]
split = json.loads((ROOT / "training/split_manifest.json").read_text())
assign_ds = json.loads((ROOT / "training/assignment_dataset_summary.json").read_text())
ref = json.loads((ROOT / "monitoring/reference_dist.json").read_text())
psi_trace = json.loads((ROOT / "monitoring/psi_trace_normal_prob.json").read_text())
normalisation = list(csv.DictReader(open(ROOT / "experiments/normalisation_experiment.csv", encoding="utf-8")))

BLUE = "17365D"
TEAL = "0F6B63"
LIGHT_BLUE = "DCE6F1"
LIGHT_TEAL = "DDEBF7"
LIGHT_GREY = "F2F2F2"
RED = "C00000"
GREEN = "008000"
GOLD = "806000"


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True, color="FFFFFF", size=font_size)
        set_cell_shading(hdr.cells[i], BLUE)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
            if ridx % 2 == 1:
                set_cell_shading(cells[i], LIGHT_GREY)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.5)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_cell_shading(table.cell(0, 0), "F7F7F7")
    p = table.cell(0, 0).paragraphs[0]
    for line in text.strip("\n").splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8)
    doc.add_paragraph()


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(2)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


def metric_row(result):
    return [
        result["variant"],
        f"{result['mean_latency_ms']:.6f}",
        f"{result['p95_latency_ms']:.6f}",
        f"{result['size_kb']:.2f}",
        f"{result['accuracy_pct']:.2f}",
        f"{result['recall_critical_pct']:.1f}",
        f"{result['energy_mj_per_inference']:.5f}",
    ]


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.65); sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(0.68); sec.right_margin = Inches(0.68)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
styles["Normal"].font.size = Pt(10)
styles["Normal"].paragraph_format.space_after = Pt(5)
styles["Normal"].paragraph_format.line_spacing = 1.08
for s in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
    styles[s].font.name = "Aptos Display"
    styles[s]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
styles["Heading 1"].font.color.rgb = RGBColor.from_string(BLUE)
styles["Heading 1"].font.size = Pt(18)
styles["Heading 2"].font.color.rgb = RGBColor.from_string(TEAL)
styles["Heading 2"].font.size = Pt(13)
styles["Heading 3"].font.color.rgb = RGBColor.from_string(GOLD)
styles["Heading 3"].font.size = Pt(11)
styles["Caption"].font.name = "Aptos"

if "Code Text" not in styles:
    st = styles.add_style("Code Text", WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = "Consolas"; st.font.size = Pt(8)

# Headers/footers
for section in doc.sections:
    hp = section.header.paragraphs[0]
    hp.text = "LogiEdge - Edge AI for FreightBridge Cold-Chain Trucks"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.size = Pt(8); hp.runs[0].font.color.rgb = RGBColor(100,100,100)
    add_page_number(section.footer.paragraphs[0])

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run("LOGIEDGE")
r.bold = True; r.font.size = Pt(34); r.font.color.rgb = RGBColor.from_string(BLUE)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run("Intelligent Edge AI Platform for Pharmaceutical Cold-Chain Logistics")
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor.from_string(TEAL)
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p3.add_run("Complete Final Assignment Report and Technical Validation Package")
r.font.size = Pt(13)
doc.add_paragraph()
cover_rows = [
    ["Course", "AIML ZG535 - Machine Learning on Edge"],
    ["Programme", "BITS Pilani Work Integrated Learning Programmes"],
    ["Business case", "FreightBridge Logistics Pvt. Ltd. - 85 refrigerated-truck pilot"],
    ["Lead contributor", "Raja Chairmapandi"],
    ["BITS ID", "2024AD05460"],
    ["Group number", "[ENTER GROUP NUMBER BEFORE SUBMISSION]"],
    ["Additional contributors", "[ENTER NAMES AND BITS IDs BEFORE SUBMISSION]"],
    ["Verification date", "16 July 2026"],
]
add_table(doc, ["Item", "Details"], cover_rows, widths=[1.7,4.7], font_size=9.5)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Recommended deployment: M3 - 35% structured pruning + Full INT8 PTQ")
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor.from_string(GREEN)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Grouped validation accuracy 99.14% | Critical recall 99.4% | Model size 3.70 KB")
r.font.size = Pt(10)
# Executive summary
h = doc.add_heading("Executive Summary", level=1)
h.paragraph_format.page_break_before = True
for text in [
    "FreightBridge's refrigerated fleet operates through rural corridors where cellular service is unavailable for 35-90 minutes at seven documented points. LogiEdge therefore keeps the safety-critical path inside the truck: sensors publish to a local Mosquitto broker, a six-feature preprocessing pipeline runs locally, a TensorFlow Lite classifier assigns Normal, Warning or Critical, and every result is written to a durable SQLite audit log before remote transmission. Cellular loss delays reporting to Pune but does not stop detection, driver alerting, chain-of-custody recording or later synchronisation.",
    "The package implements every task from Modules 1-6: quantitative constraint analysis, hardware and Roofline selection, sensor simulation, MQTT integration, frozen-statistics preprocessing, feature-level fusion, model training, full INT8 conversion, structured pruning, Docker layer caching, a seven-task idempotent Ansible playbook, PSI monitoring and five-metric model benchmarking. The exact minimum Task D1 run produces 117 Normal, 87 Warning and 87 Critical windows; the final experiment expands this to ten independent simulated trucks so that two complete trucks form a leakage-safe 20% holdout.",
    "The final models exceed both mandatory gates. M1 reaches 99.48% grouped accuracy, M2 reaches 99.31%, and M3 reaches 99.14%. All three provide 99.4% Critical recall, above the required 95%. Development-host latency is approximately 1.4 microseconds per window for all variants and is below measurement-separation limits; the operational decision is therefore governed by safety recall, size and deployment efficiency. M3 is recommended because it is the smallest model at 3.70 KB while preserving the same Critical recall.",
    "PSI is monitored on the Normal-class output confidence, interpreted as a cargo-health score. The rolling monitor crosses 0.25 three minutes after a combined fault is injected, reaches 2.614 and returns to 0.037 after clean data replaces the anomalous samples. Automated verification completed successfully with Python compilation and eight passing tests. Live Docker, Mosquitto and Ansible screen evidence must still be recorded on the student's target machine because those services are unavailable in this report-build environment; the exact commands and expected outputs are included in the package.",
]:
    doc.add_paragraph(text)

add_table(doc, ["Mandatory gate", "Requirement", "Verified result", "Status"], [
    ["Validation accuracy", "> 88%", "M3 = 99.14%", "PASS"],
    ["Critical recall", "> 95%", "M3 = 99.4%", "PASS"],
    ["PSI alert", "> 0.25 within 5 min", "0.390 at 3 min; max 2.614", "PASS"],
    ["PSI recovery", "< 0.10", "0.037", "PASS"],
    ["Ansible structure", "Exactly 7 tasks", "7 top-level tasks", "PASS"],
    ["Automated tests", "No regression failures", "8 passed", "PASS"],
], widths=[1.3,1.5,2.0,0.8], font_size=8.5)

doc.add_page_break()

# A
A1_PARAS = [
"Latency. A refrigeration failure can increase cargo temperature by 1 degree C per minute, so a 90-second delay permits a 1.5 degree C rise. The complete response time is sampling and window formation plus preprocessing, inference, decision and alerting. LogiEdge uses a 30-second window and ten-second step, so the worst alignment adds about 40 seconds before a new feature vector is available; local inference and decision are negligible in comparison. Under usable rural coverage, a conservative 150-500 ms cellular round trip could permit cloud inference in a few seconds, but retransmission, handover and broker delay are uncontrolled. More importantly, a 35-90 minute outage makes cloud response time effectively infinite. Cloud inference is therefore sometimes fast but cannot guarantee the 90-second safety SLA. Edge inference can guarantee a local decision while the modem is disconnected.",
"Bandwidth. Assuming float32 samples, temperature produces 1 x 4 x 86,400 = 345,600 bytes or 0.346 MB per truck per day. Raw three-axis vibration produces 500 x 3 x 4 x 86,400 = 518.4 MB per day. Approximately 20 compact door events add less than 0.001 MB, giving about 518.7 MB per truck per day before protocol overhead. At Rs 0.10 per MB this costs Rs 51.87 per truck per day, about Rs 1.61 million per year for 85 trucks and Rs 5.02 million for 265 trucks. LogiEdge transmits only five-minute 200-byte heartbeats and up to twenty 300-byte alert records, approximately 63.6 KB per truck per day. This is an 8,156-fold reduction and protects the M2M link from raw vibration traffic.",
"Connectivity. During the seven Nashik-Aurangabad coverage gaps, a cloud-only classifier receives no current sensor values, produces no warning, and leaves an incomplete chain-of-custody record unless a separate local buffer has already been engineered. LogiEdge keeps the local broker, preprocessing, TFLite inference, driver alert and SQLite audit log active. Each event has a unique ID and independent inference and alert acknowledgement flags. The cellular uplink retries with QoS 1, and after reconnection the oldest pending rows are replayed until the backlog is empty. Operations-centre visibility is delayed, but the safety action and evidence are not lost.",
"Privacy. On-device inference means continuous high-frequency raw vibration and temperature streams do not leave the truck. Only alerts, health summaries and authorised audit records are transmitted, reducing third-party exposure and supporting the client's contractual requirement. Privacy is completed by per-truck credentials, TLS on the uplink, encrypted local storage, role-based access, signed model packages, retention rules and auditable access. Edge inference alone is not claimed to guarantee privacy; it reduces the data surface and makes technical controls enforceable even when cloud connectivity is absent."
]
a1_words = len(" ".join(A1_PARAS).replace("-", " ").split())
doc.add_heading("A. System Architecture and Deployment Justification", level=1)
doc.add_heading(f"Task A1 - Constraint Analysis ({a1_words} words)", level=2)
for para in A1_PARAS:
    doc.add_paragraph(para)

doc.add_heading("Task A2 - Complete System Architecture", level=2)
doc.add_picture(str(ROOT / "scenario_architecture/system_architecture.png"), width=Inches(7.0))
add_caption(doc, "Figure A1. Per-truck edge architecture. The complete sensor-to-alert path remains inside the dashed truck boundary.")
doc.add_paragraph("The architecture separates the always-available local sensor path from the failure-prone cellular uplink. Temperature, compressor vibration and door events enter a localhost Mosquitto broker. The edge service performs filtering, windowing, normalisation and TFLite inference, records the complete probability vector and context in SQLite, issues local alerts and transmits through a separate uplink client. The operations centre receives alert escalation, chain-of-custody records, PSI telemetry and OTA deployment commands. The model update path is deliberately reverse-direction and does not participate in real-time inference.")

doc.add_page_break()

# B

doc.add_heading("B. Constraint Triangle and Roofline Analysis", level=1)
doc.add_heading("Task B1 - Constraint Triangle Application", level=2)
add_table(doc, ["Option", "Unit cost", "85-truck pilot", "265-truck scale", "Power", "Dominant weakness"], [
    ["Raspberry Pi 5 8 GB + Hailo-8L", "Rs 15,000", "Rs 12.75 lakh", "Rs 39.75 lakh", "7.5 W", "Balanced; requires automotive enclosure and power conditioning"],
    ["Jetson Orin Nano Super", "Rs 45,000", "Rs 38.25 lakh", "Rs 119.25 lakh", "15 W", "Exceeds 10 W budget and is over-specified"],
    ["STM32H7 custom MCU", "Rs 3,500", "Rs 2.98 lakh", "Rs 9.28 lakh", "0.4 W", "Memory, development and MLOps limitations"],
], widths=[1.35,0.75,0.9,0.95,0.55,2.0], font_size=8)
doc.add_picture(str(FIG / "hardware_cost_power.png"), width=Inches(6.8))
add_caption(doc, "Figure B1. Fleet cost and power evidence for the Constraint Triangle.")
doc.add_paragraph("The dominant vertex is power-constrained and cost-sensitive deployment subject to a hard latency floor. The 90-second SLA does not justify 67 TOPS: all candidates must first prove safe response, after which power, fleet cost and maintainability decide. Raspberry Pi 5 plus AI HAT+ is selected because the stated 7.5 W remains below the 10 W AI budget, Linux supports Mosquitto, Docker, SQLite, Ansible and secure OTA, and the 85-truck hardware cost is one-third of Jetson. Jetson is rejected because 15 W exceeds the budget and Rs 1.19 crore at full scale is disproportionate to a 45 MFLOP model. STM32H7 is rejected for the pilot because the 18 MB working set, external-memory dependence, lack of Docker and substantial firmware effort create schedule and maintenance risk. It remains a future cost-down option after the model and field signals stabilise.")

doc.add_heading("Task B2 - Arithmetic Intensity and Roofline", level=2)
add_table(doc, ["Quantity", "Calculation", "Result"], [
    ["Arithmetic intensity", "45 x 10^6 FLOP / 18 x 10^6 byte", "2.50 FLOP/byte"],
    ["Ridge point", "16 GFLOP/s / 12 GB/s", "1.333 FLOP/byte"],
    ["Classification", "2.50 > 1.333", "Compute-bound"],
    ["Compute time", "45 MFLOP / 16 GFLOP/s", "2.8125 ms"],
    ["Memory time", "18 MB / 12 GB/s", "1.500 ms"],
    ["Roofline latency", "max(compute, memory)", "2.81 ms theoretical"],
], widths=[1.35,3.0,1.6], font_size=8.7)
doc.add_paragraph("Because arithmetic intensity lies to the right of the ridge point, the processor reaches its compute roof before saturating memory bandwidth. The primary Roofline action is therefore to reduce operations or increase effective compute throughput: full INT8 quantisation, physical removal of hidden units, operator fusion, NEON acceleration or Hailo offload. Memory locality still affects real latency and energy, but it is not the theoretical bottleneck under the supplied parameters. The 2.81 ms estimate is 32,000 times faster than 90 seconds and excludes software overhead, so target-device end-to-end measurements remain necessary.")

# C

doc.add_heading("C. Sensor Simulation, Preprocessing and Data Fusion", level=1)
doc.add_heading("Task C1 - Sensor Simulator", level=2)
add_table(doc, ["Stream", "Frequency", "Normal model", "Fault model", "MQTT topic suffix"], [
    ["Temperature", "1 Hz", "N(4.0 C, 0.3)", "+0.08 C per reading", "/sensors/temperature"],
    ["Vibration RMS", "0.5 Hz", "N(0.45 g, 0.05)", "N(1.2 g, 0.15)", "/sensors/vibration"],
    ["Door", "Discrete", "OPEN/CLOSE timestamp", "Context during all modes", "/sensors/door"],
], widths=[1.1,0.7,1.4,1.5,1.6], font_size=8.5)
add_code(doc, """python data_pipeline/simulator.py --anomaly none --duration 600
python data_pipeline/simulator.py --anomaly temp_drift --duration 900
python data_pipeline/simulator.py --anomaly vibration --duration 900
python data_pipeline/simulator.py --anomaly combined --duration 900""")
doc.add_paragraph("The generator uses one implementation for both live MQTT publication and offline dataset generation. Vibration at 0.5 Hz correctly means one RMS message every two seconds. Temperature drift is accelerated synthetic degradation; Warning runs are bounded within the warning region, while combined failure is allowed to breach the Critical region. Door events are preserved for the audit trail and local rule context but are not silently added as a seventh feature because Task C2 fixes the model input at six values.")

doc.add_heading("Task C2 - Complete Preprocessing Pipeline", level=2)
add_table(doc, ["Order", "Operation", "Implementation"], [
    ["1", "Filtering", "Causal five-sample moving average on temperature and vibration"],
    ["2", "Windowing", "30-second window, ten-second step, 20-second overlap"],
    ["3", "Temperature features", "Mean, standard deviation, least-squares rate of change in C/min"],
    ["4", "Vibration features", "Window RMS, absolute peak, Fisher excess kurtosis"],
    ["5", "Normalisation", "z = (x - mean) / std using frozen training_stats.npy"],
    ["6", "Inference", "Six-value float vector quantised according to TFLite input scale and zero-point"],
], widths=[0.55,1.35,4.4], font_size=8.5)
doc.add_picture(str(FIG / "normalisation_sensitivity.png"), width=Inches(6.5))
add_caption(doc, "Figure C1. Mandatory correct-statistics versus 3-sigma shifted-statistics experiment.")
add_table(doc, ["Normalisation", "Accuracy", "Normal recall", "Warning recall", "Critical recall", "Critical called Normal"], [
    [r["normalisation"], r["accuracy_pct"]+"%", r["recall_normal_pct"]+"%", r["recall_warning_pct"]+"%", r["recall_critical_pct"]+"%", r["critical_windows_called_normal"]]
    for r in normalisation
], widths=[1.4,0.8,0.85,0.9,0.9,1.15], font_size=8)
doc.add_paragraph("The experiment demonstrates why runtime statistics must never adapt to live cargo. A +3 sigma corruption happened not to reduce overall accuracy on this synthetic holdout, but it changed class behaviour; a -3 sigma shift reduced accuracy to 68.90%. The result is reported honestly rather than forcing a deterioration in only one direction. The engineering conclusion remains that commissioning statistics must be versioned with the model and changed only through controlled validation.")

doc.add_heading("Task C3 - Feature-Level Fusion Justification", level=2)
doc.add_paragraph("Feature-level fusion is the best fit because temperature and vibration have different rates, units and physical meanings. Each stream is filtered and summarised using sensor-specific statistics, then the six features are concatenated before classification. Raw data-level fusion would require artificial resampling of 1 Hz temperature, 0.5 Hz RMS vibration and asynchronous door events, increasing memory and potentially distorting kurtosis. Decision-level fusion would require separate models and conflict logic and would weaken the interaction between temperature rate of change and bearing-wear vibration that distinguishes Warning from Critical. Door state remains contextual metadata and can enrich driver guidance without violating the required six-feature contract.")

# D

doc.add_heading("D. Dataset, Model Lifecycle, Docker and Pipeline Mapping", level=1)
doc.add_heading("Task D1 - Dataset Generation and Model Training", level=2)
add_table(doc, ["Class", "Mode", "Required duration", "Exact run windows", "Expanded grouped windows"], [
    ["0 Normal", "none", "20 min", "117", str(split["train_class_counts"].get("0",0)+split["validation_class_counts"].get("0",0))],
    ["1 Warning", "temp_drift", "15 min", "87", str(split["train_class_counts"].get("1",0)+split["validation_class_counts"].get("1",0))],
    ["2 Critical", "combined", "15 min", "87", str(split["train_class_counts"].get("2",0)+split["validation_class_counts"].get("2",0))],
], widths=[1.0,1.0,1.1,1.3,1.55], font_size=8.5)
doc.add_picture(str(FIG / "dataset_distribution.png"), width=Inches(6.6))
add_caption(doc, "Figure D1. Exact assignment run and final grouped train/validation composition.")
doc.add_paragraph("The prescribed minimum run is retained in assignment_dataset.npz. The rounded 120/90/90 expectations become 117/87/87 because the first complete window closes at 30 seconds and the simulator emits through duration minus one second. For final training, the same three durations are repeated for ten independent simulated trucks. Eight truck groups provide 2,328 windows and two unseen truck groups provide 582 windows. No truck group appears in both partitions. This avoids optimistic leakage from adjacent windows that share 20 of 30 seconds.")
add_table(doc, ["Variant", "Accuracy", "Normal recall", "Warning recall", "Critical recall", "Parameters / size"], [
    ["M1 FP32", f"{m1['val_accuracy']*100:.2f}%", f"{m1['recall']['Normal']*100:.1f}%", f"{m1['recall']['Warning']*100:.1f}%", f"{m1['recall']['Critical']*100:.1f}%", f"{m1['params']} / {m1['tflite_bytes']/1024:.2f} KB"],
    ["M2 Full INT8", f"{m2['accuracy']*100:.2f}%", f"{m2['recall']['Normal']*100:.1f}%", f"{m2['recall']['Warning']*100:.1f}%", f"{m2['recall']['Critical']*100:.1f}%", f"{m2['size_kb']:.2f} KB"],
    ["M3 Pruned INT8", f"{m3['accuracy']*100:.2f}%", f"{m3['recall']['Normal']*100:.1f}%", f"{m3['recall']['Warning']*100:.1f}%", f"{m3['recall']['Critical']*100:.1f}%", f"{m3['params']} / {m3['size_kb']:.2f} KB"],
], widths=[1.2,0.8,0.9,0.9,0.9,1.25], font_size=8.2)
doc.add_picture(str(FIG / "confusion_matrices.png"), width=Inches(6.8))
add_caption(doc, "Figure D2. Confusion matrices on the complete-truck grouped holdout.")
doc.add_paragraph("The baseline is a six-input MLP with Dense(32, ReLU), Dense(16, ReLU) and Dense(3, softmax). Its 99.48% grouped validation accuracy exceeds the mandatory 88% gate. Critical recall is 99.4%, and the confusion matrix shows one Critical window classified as Warning and none classified as Normal.")

doc.add_heading("Task D2 - Docker Containerisation and OTA Layer Cache", level=2)
add_code(doc, """FROM python:3.11-slim
COPY requirements.txt /app/requirements.txt
RUN pip install --no-cache-dir -r /app/requirements.txt
COPY preprocessing.py tflite_eval.py psi.py inference_service.py /app/
COPY training_stats.npy reference_dist.json /data/
COPY model.tflite /data/model.tflite
ENV MODEL_PATH=/data/model.tflite
CMD [\"python\", \"-u\", \"/app/inference_service.py\"]""")
doc.add_paragraph("All package installation and immutable application-code layers precede the model layer. MODEL_PATH is read at runtime, and the inference result topic resolves exactly to logibridge/trucks/{truck_id}/inference. A model-only rebuild reuses the base, dependency, code and configuration layers. Using the demo assumption of a 150 MB complete image, distributing the entire image to 85 trucks would transfer 12,750 MB and cost Rs 1,275. The 3.70 KB M3 layer transfers approximately 0.307 MB for the fleet and costs about Rs 0.03, saving approximately 99.998%. The demo script captures the initial build and second build with CACHED shown for unchanged layers.")

doc.add_heading("Task D3 - Ten-Stage Edge ML Pipeline Mapping", level=2)
ten = [
("1. Business problem", "Prevent pharmaceutical spoilage and documentation failure by detecting temperature and compressor anomalies within 90 seconds even without cellular service."),
("2. Sensor acquisition", "The truck publishes temperature at 1 Hz, compressor RMS vibration at 0.5 Hz and timestamped door edges to its local broker."),
("3. Data validation and filtering", "Payloads are schema-checked by field use, and each continuous stream passes through a causal five-sample moving average."),
("4. Windowing and feature engineering", "A 30-second window advances every ten seconds and produces temperature mean, standard deviation, rate of change, vibration RMS, peak and kurtosis."),
("5. Normalisation and labelling", "Features use frozen clean commissioning statistics; simulator modes and fault onset create Normal, Warning and Critical labels."),
("6. Model training", "A compact 32-16 MLP is trained on eight complete simulated truck groups with early stopping and learning-rate reduction."),
("7. Evaluation", "Two unseen truck groups form the 20% holdout; accuracy, confusion matrices and class recalls enforce the 88% and 95% gates."),
("8. Optimisation", "The baseline is converted to full INT8 with 250 calibration samples; M3 uses PolynomialDecay pruning, physical unit removal and full INT8 PTQ."),
("9. Packaging and deployment", "Preprocessing, TFLite inference, PSI and durable logging are packaged in Python 3.11 slim and deployed with an exactly seven-task Ansible playbook."),
("10. Monitoring and update", "Rolling PSI, container health, backlog state and fleet alerts support a six-week canary OTA cycle with rollback."),
]
for title, body in ten:
    p = doc.add_paragraph()
    r = p.add_run(title + ". "); r.bold = True; r.font.color.rgb = RGBColor.from_string(TEAL)
    p.add_run(body)

doc.add_page_break()

# E

doc.add_heading("E. Edge MLOps, Ansible and OTA Strategy", level=1)
doc.add_heading("Task E1 - PSI Drift Monitoring", level=2)
doc.add_paragraph("The monitored scalar is the model's Normal-class output confidence p(Normal), interpreted as a cargo-health confidence score. This choice remains an output confidence score while avoiding a known failure of maximum softmax confidence: a competent classifier can be highly confident for both healthy and Critical cargo, so max confidence may not move between the required bins. The reference uses 300 clean Normal windows from ten truck biases, the live monitor retains the latest 100 scores and evaluates every 60 seconds using bins [0,0.25), [0.25,0.50), [0.50,0.75) and [0.75,1.0]. Empty bins use a small numerical floor to prevent log division by zero.")
add_table(doc, ["PSI requirement", "Verified result", "Status"], [
    ["Clean reference", f"300 windows; distribution {ref['distribution']}", "PASS"],
    ["Alert threshold", "Crossed 0.25 at 3.0 minutes after injection", "PASS"],
    ["Combined-fault maximum", "PSI = 2.614", "PASS"],
    ["Recovery threshold", "Final PSI = 0.037 < 0.10", "PASS"],
], widths=[1.9,3.5,0.8], font_size=8.5)
doc.add_picture(str(FIG / "psi_drift_recovery.png"), width=Inches(6.6))
add_caption(doc, "Figure E1. PSI crosses the drift threshold within five minutes and later clears below 0.10.")
add_code(doc, """[LOGIBRIDGE DRIFT ALERT] PSI=0.390
...
[LOGIBRIDGE] drift cleared, PSI=0.039 < 0.10""")

doc.add_heading("Task E2 - Exactly Seven Ansible Tasks", level=2)
ansible_rows = [
["1", "Create /opt/logibridge/data"],
["2", "Copy new model.tflite"],
["3", "Copy reference_dist.json and frozen training_stats.npy in one loop task"],
["4", "Stop the container only when model or configuration changed"],
["5", "Pull the versioned image from the local registry"],
["6", "Start the container with MODEL_PATH, truck ID, local broker and uplink variables"],
["7", "Wait up to 15 seconds and fail unless the container is running"],
]
add_table(doc, ["Task", "Implementation"], ansible_rows, widths=[0.55,5.7], font_size=8.5)
doc.add_paragraph("Idempotency is achieved by state-based community.docker modules and by stopping only when copied content changes. The final information task is explicitly changed_when: false. Static tests verify exactly seven top-level tasks and the corrected /data mount. The first and second live playbook runs must be recorded on a Docker-and-Ansible-capable machine; the second recap is expected to show changed=0.")

doc.add_heading("Task E3 - OTA Strategy Selection", level=2)
add_table(doc, ["Strategy", "Model transfer", "Cost at Rs 0.10/MB", "Risk interpretation"], [
    ["Full replacement", "23.242 MB to 85 trucks", "Rs 2.324", "Fast but exposes all 85 trucks at once"],
    ["Canary 10 then 75", "2.734 MB first; 23.242 MB total", "Rs 0.273 first; Rs 2.324 total", "Limits initial safety blast radius"],
    ["Shadow mode", "23.242 MB model + 205.63 MB for 7-day dual telemetry", "Approx. Rs 22.89", "Safest comparison but more compute, telemetry and operational complexity"],
], widths=[1.15,1.7,1.4,2.25], font_size=8)
doc.add_paragraph("The shadow estimate assumes one additional 40-byte result every ten seconds for seven days across 85 trucks. Canary is recommended because cold-chain safety risk, not bandwidth cost, dominates. Ten trucks should represent different routes, refrigeration units and connectivity conditions. Promotion criteria are Critical recall, false-alert rate, PSI, inference-service uptime, backlog replay and incident review. Full replacement is rejected because one defective model reaches the whole fleet. Full-fleet shadow is rejected as the default because it doubles inference and telemetry complexity, although shadow comparison can be used inside the ten-truck canary before activation.")

# F

doc.add_heading("F. Model Optimisation, Benchmarking and Deployment Decision", level=1)
doc.add_heading("Task F1 - Three Model Variants", level=2)
add_table(doc, ["Variant", "Method", "Verification"], [
    ["M1 FP32", "Standard Keras 32-16 MLP and FP32 TFLite export", f"803 parameters; {m1['tflite_bytes']/1024:.2f} KB"],
    ["M2 PTQ INT8", "Full integer PTQ with 250 training-only representative samples", "INT8 input/output and 4.46 KB"],
    ["M3 prune + INT8", "TFMOT PolynomialDecay to 35%, strip pruning, physically retain 21 and 10 hidden units, fine-tune, Full INT8", f"400 parameters; {m3['size_kb']:.2f} KB"],
], widths=[1.15,3.8,1.55], font_size=8)
doc.add_paragraph("The phrase structured filter pruning is adapted correctly for a dense MLP: TFMOT first applies 35% magnitude pruning under PolynomialDecay, then complete hidden units are selected by column norm and physically removed. The architecture changes from 32 and 16 units to 21 and 10 units. This produces a genuinely smaller dense network before INT8 conversion rather than merely storing unstructured zeros.")

doc.add_heading("Task F2 - Five-Metric Benchmark", level=2)
add_table(doc, ["Variant", "Mean latency ms", "p95 ms", "Size KB", "Accuracy %", "Critical recall %", "Energy mJ"], [metric_row(x) for x in bench], widths=[1.3,1.0,0.8,0.7,0.8,0.95,0.8], font_size=7.7)
doc.add_picture(str(ROOT / "optimisation/results/pareto_chart.png"), width=Inches(6.9))
add_caption(doc, "Figure F1. Pareto analysis. Latency differences are smaller than development-host measurement noise.")
doc.add_paragraph("Each variant uses the same LiteRT interpreter, one-window batch, validation set and host configuration. Ten warm-up calls are excluded, followed by 200 measured runs; the protocol is repeated seven times. Model size is the actual file length. Energy is the required comparative estimate E = P x t using a 15 W laptop TDP and CPU utilisation sampled over a longer identical invocation loop. It is not a direct electrical measurement. The latest run gives approximately 1.41-1.42 microseconds for all variants, so no honest speed ranking is claimed on this x86_64 host.")

doc.add_heading("Task F3 - Deployment Recommendation to the Operations Director", level=2)
doc.add_paragraph("Deploy M3 to the 85 refrigerated trucks through a ten-truck canary. The end-to-end 90-second SLA includes approximately 40 seconds for window formation in the worst alignment, less than one second for preprocessing and local decision, and a large safety reserve. M3 inference is roughly 0.00142 ms on the development host, so inference consumes a negligible fraction of the budget. On the selected Raspberry Pi 5, the 3.70 KB model and small activation tensors are trivial relative to 8 GB LPDDR4X memory and normal persistent storage; the operational memory load is dominated by Python, LiteRT, Mosquitto, Docker and SQLite rather than model weights.")
doc.add_paragraph("M3 achieves 99.14% grouped accuracy and 99.4% Critical recall, exceeding the mandatory 95% safety floor. It is 29% smaller than the FP32 TFLite file and reduces the network from 803 to 400 parameters without a Critical-recall loss. M1 has 0.34 percentage points more accuracy, but this does not justify the larger OTA artefact when the safety metric is equal. M2 is a sound fallback, but it is larger than M3 with no recall advantage. Final approval requires canary evidence on actual Raspberry Pi hardware, including MQTT-to-alert latency, power, temperature, throttling and cellular backlog recovery.")

# Security and limitations

doc.add_heading("G. Security, Chain of Custody and Operational Controls", level=1)
for item in [
    "Durable-first audit: every inference is inserted into SQLite WAL storage before remote publication.",
    "Unique event identity and independent acknowledgement flags prevent duplicate loss and distinguish inference synchronisation from alert synchronisation.",
    "Per-truck credentials and TLS protect the uplink; localhost sensor traffic remains inside the truck.",
    "Model and reference files should be signed and hash-verified before container restart.",
    "Role-based access, audit logging and retention rules support pharmaceutical client contracts.",
    "Future production hardening should add secure time, hash chaining, encrypted storage, schema validation and certificate rotation.",
]:
    add_bullet(doc, item)

doc.add_heading("H. Limitations and Final Validation Gates", level=1)
add_table(doc, ["Area", "Current evidence", "Required production / submission action"], [
    ["Data realism", "Physics-informed synthetic simulator", "Validate with labelled reefer sensor data and vibration-only failures"],
    ["Hardware", "Roofline analysis and x86 benchmark", "Benchmark on Raspberry Pi 5 + AI HAT+, including power and thermal throttling"],
    ["Docker", "Dockerfile and cache demo script", "Record initial and model-only rebuild showing cached layers"],
    ["Ansible", "Exactly seven tasks and static tests", "Record two runs; second recap must show changed=0"],
    ["Connectivity", "Two-broker architecture and durable-log tests", "Record uplink broker stop/start and oldest-first replay"],
    ["Submission metadata", "Lead contributor populated", "Enter group number, all members, repository URL and video URL"],
], widths=[1.1,2.2,3.25], font_size=8.2)

# Conclusion

doc.add_heading("Conclusion", level=1)
doc.add_paragraph("LogiEdge meets the core purpose of Machine Learning on Edge: a safety decision remains available where the data is produced even when the wide-area network is absent. The package provides quantitative deployment justification, a realistic multi-stream simulator, six-feature fusion, frozen commissioning statistics, leakage-safe validation, full INT8 deployment, structured model reduction, durable store-and-forward, PSI monitoring, idempotent fleet deployment logic and a defensible model recommendation. M3 is selected for the 85-truck pilot because it combines 99.14% grouped validation accuracy, 99.4% Critical recall and a 3.70 KB deployment artefact. The remaining steps are field verification and the student's live demo recording, not additional synthetic model tuning.")

# References

doc.add_heading("References", level=1)
refs = [
    "BITS Pilani WILP. AIML ZG535 Machine Learning on Edge - LogiEdge mini-project problem statement and module notes.",
    "TensorFlow. TensorFlow Lite / LiteRT post-training integer quantisation documentation.",
    "TensorFlow Model Optimization Toolkit. Pruning with Keras and PolynomialDecay documentation.",
    "Eclipse Foundation. Mosquitto and MQTT QoS documentation.",
    "Docker. Dockerfile layer caching and content-addressed image documentation.",
    "Ansible. community.docker collection documentation.",
    "Raspberry Pi. Raspberry Pi 5 and AI HAT+ technical documentation.",
]
for reftext in refs:
    add_number(doc, reftext)

# Appendix

doc.add_heading("Appendix A - Submission File Index", level=1)
files = [
("reports/LogiEdge_Complete_Final_Report.pdf", "Final PDF report for LMS submission"),
("reports/LogiEdge_Complete_Final_Report.docx", "Editable report source"),
("training/assignment_dataset.npz", "Exact Task D1 minimum-duration dataset"),
("training/dataset.npz", "Expanded grouped dataset with train/validation indices"),
("training/models/", "M1, M2 and M3 models and metrics"),
("data_pipeline/simulator.py", "Three-stream CLI simulator and MQTT publisher"),
("data_pipeline/preprocessing.py", "Moving average, sliding windows, six features and frozen normaliser"),
("inference/", "Python 3.11 Docker inference service"),
("monitoring/", "Reference distribution, PSI library, drift and recovery trace"),
("deployment/logibridge_deploy.yml", "Exactly seven-task Ansible deployment"),
("optimisation/results/", "Five-metric benchmark and Pareto chart"),
("evidence/verification_run_2026-07-16.log", "Executed local verification log"),
("demo/demo_script.md", "15-20 minute demo sequence and expected evidence"),
("reports/ASSIGNMENT_COMPLIANCE_MATRIX.md", "Requirement-by-requirement audit"),
]
add_table(doc, ["File / folder", "Purpose"], files, widths=[2.8,3.6], font_size=8.2)

doc.add_heading("Appendix B - End-to-End Execution", level=1)
add_code(doc, r"""# Windows PowerShell
.\setup_training_env.ps1
.\.venv311\Scripts\Activate.ps1
python training\generate_assignment_dataset.py
python training\generate_dataset.py
python training\train_model.py
python training\convert_ptq.py
python training\prune_quantise.py
python experiments\normalisation_experiment.py
python optimisation\benchmark.py
python monitoring\drift_monitor.py --mode reference --score normal_prob
python monitoring\drift_monitor.py --mode simulate --score normal_prob
pytest -q""")
add_code(doc, """# Docker/MQTT demo
cd demo
docker compose up -d
# start the inference service or container, then run:
python ../data_pipeline/simulator.py --anomaly combined --duration 600 --speed 20
docker compose stop uplink-broker
docker compose start uplink-broker
# follow demo/demo_script.md for cache and Ansible evidence""")

doc.add_heading("Appendix C - Final Submission Checklist", level=1)
for item in [
    "Replace the group-number and contributor placeholders on the cover.",
    "Insert the private GitHub repository URL and grant the instructor collaborator access.",
    "Run the Docker layer-cache demonstration and retain the first and second build output.",
    "Run logibridge_deploy.yml twice and capture the second PLAY RECAP with changed=0.",
    "Demonstrate local inference while only the uplink broker is stopped, then show backlog replay after restart.",
    "Record PSI clean, combined-fault alert and recovery output.",
    "Upload the 15-20 minute demo and replace demo/demo_video_link.txt with the accessible URL.",
    "Rename the final PDF to the exact LMS naming convention if a group-number prefix is required.",
]:
    add_bullet(doc, item)


# Save
out = REPORTS / "LogiEdge_Complete_Final_Report.docx"
doc.save(out)
print(f"A1 word count: {a1_words}")
print(f"Saved {out}")
